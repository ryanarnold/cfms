from django.shortcuts import render
from django.core.urlresolvers import reverse
from django.http import HttpResponseRedirect, HttpResponse
from docx import Document

# Create your views here.
def reports(request):
	return render(request, 'reports/reports.html', {})

def monthly_detail(request, report_type):
	document = Document()
	paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
	document.add_heading('REPORT')
	document.save('test.docx')
	# return HttpResponse(document.save('test.docx'))
	
	response = HttpResponse(content_type='application/docx')
	response['Content-Disposition'] = 'attachment; filename="test.docx"'
	f = open('test.docx', 'r')
	return response

	# return HttpResponseRedirect(reverse('reports'))